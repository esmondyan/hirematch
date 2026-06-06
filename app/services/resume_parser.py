import io
from pathlib import Path

import fitz  # PyMuPDF
import pymupdf4llm
from docx import Document

MAX_SIZE_BYTES = 5 * 1024 * 1024  # 5MB

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class ResumeParserError(Exception):
    pass


def parse_resume(file_bytes: bytes, filename: str, format: str = "markdown") -> str:
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ResumeParserError(f"不支持的文件格式: {ext}，仅支持 PDF/DOCX/TXT")

    if len(file_bytes) > MAX_SIZE_BYTES:
        raise ResumeParserError(f"文件过大: {len(file_bytes) / 1024 / 1024:.1f}MB，限制 {MAX_SIZE_BYTES / 1024 / 1024:.0f}MB")

    if ext == ".pdf":
        return _parse_pdf_with_ocr_fallback(file_bytes, format)
    elif ext == ".docx":
        text = _parse_docx(file_bytes)
        return _to_basic_markdown(text) if format == "markdown" else text
    else:
        text = _parse_txt(file_bytes)
        return _to_basic_markdown(text) if format == "markdown" else text


def _parse_pdf_text(file_bytes: bytes) -> str:
    """Plain text extraction via PyMuPDF (fast, no layout analysis)."""
    text_parts = []
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    for page in doc:
        page_text = page.get_text()
        if page_text and page_text.strip():
            text_parts.append(page_text.strip())
    doc.close()
    if not text_parts:
        raise ResumeParserError("PDF无法提取文本，可能为扫描件或图片PDF")
    return "\n\n".join(text_parts)


def _parse_pdf_ocr(file_bytes: bytes) -> str:
    """OCR fallback for scanned/image-based PDFs using PyMuPDF's built-in OCR.

    Requires Tesseract OCR and language data to be installed on the system:
      - Windows: install Tesseract from https://github.com/UB-Mannheim/tesseract/wiki
      - Linux:   apt-get install tesseract-ocr tesseract-ocr-chi-sim tesseract-ocr-eng
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts = []
    try:
        for page in doc:
            try:
                tp = page.get_textpage_ocr(flags=3, language='chi_sim+eng', dpi=300)
                page_text = page.get_text(textpage=tp)
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
            except Exception:
                continue  # Skip pages where OCR fails
    except Exception as e:
        raise ResumeParserError(
            f"PDF OCR 失败（请确保已安装 Tesseract OCR 及中文语言包 chi_sim）: {e}"
        )
    finally:
        doc.close()

    if not text_parts:
        raise ResumeParserError("PDF无法提取文本，OCR 也未识别到文字（可能为纯图片或扫描质量过低）")
    return "\n\n".join(text_parts)


def _parse_pdf_with_ocr_fallback(file_bytes: bytes, format: str) -> str:
    """Try standard extraction first, fall back to OCR for scanned/image PDFs."""
    # First attempt: standard text/markdown extraction
    try:
        result = _parse_pdf_markdown(file_bytes) if format == "markdown" else _parse_pdf_text(file_bytes)
        # If we got substantial text, return it
        if len(result.strip()) >= 50:
            return result
        # Very little text extracted — likely a scanned/image PDF, try OCR
    except ResumeParserError:
        pass  # Standard extraction failed, try OCR

    # OCR fallback
    try:
        ocr_result = _parse_pdf_ocr(file_bytes)
        if format == "markdown":
            return _to_basic_markdown(ocr_result)
        return ocr_result
    except ResumeParserError:
        raise  # Both standard and OCR failed


def _parse_pdf_markdown(file_bytes: bytes) -> str:
    """Structured markdown extraction via pymupdf4llm (layout-aware)."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        md_text = pymupdf4llm.to_markdown(doc)
    finally:
        doc.close()
    if not md_text or not md_text.strip():
        raise ResumeParserError("PDF无法提取文本，可能为扫描件或图片PDF")
    return md_text.strip()


def _parse_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
            if row_text.strip():
                text_parts.append(row_text)
    if not text_parts:
        raise ResumeParserError("DOCX文件内容为空")
    return "\n".join(text_parts)


def _parse_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise ResumeParserError("TXT文件编码无法识别")


def reparse_resume_from_file(file_path: str, format: str = "markdown") -> str:
    """Re-parse a previously saved resume file from disk.

    Uses the current parser (which may have been upgraded since the
    original upload), so re-analysis always benefits from parser improvements.
    """
    import os
    if not os.path.exists(file_path):
        raise ResumeParserError(f"原始简历文件不存在: {file_path}")
    with open(file_path, "rb") as f:
        file_bytes = f.read()
    filename = os.path.basename(file_path)
    # Remove the "original_" prefix added during upload
    if filename.startswith("original_"):
        filename = filename[len("original_"):]
    return parse_resume(file_bytes, filename, format=format)


def _to_basic_markdown(text: str) -> str:
    """Convert plain text to basic markdown format for consistency.

    Heuristic: blank-line-separated blocks become paragraphs,
    lines that look like headers get # prefix.
    """
    lines = text.split('\n')
    result = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            result.append('')
            continue
        # Short line, no ending punctuation, might be a header
        if len(stripped) <= 20 and not stripped.endswith(('。', '，', '、', '.', ',', ';', '；')):
            # Check if it looks like a section header (common in resumes)
            if any(kw in stripped for kw in ['经历', '背景', '技能', '项目', '评价', '信息', '意向', '证书', '荣誉']):
                result.append(f'## {stripped}')
                continue
        result.append(stripped)
    return '\n\n'.join(result)
